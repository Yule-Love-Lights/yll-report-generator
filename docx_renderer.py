"""Renders the structured report JSON (from Claude) into a polished .docx
with real, clickable Word checkboxes on action items."""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import io


def _add_checkbox_paragraph(doc, text, style_size=11):
    """Adds a paragraph with a real interactive Word checkbox content
    control, followed by the item text. Toggles when clicked in Word."""
    p = doc.add_paragraph()

    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')

    id_el = OxmlElement('w:id')
    id_el.set(qn('w:val'), str(abs(hash(text)) % 1000000000))
    sdtPr.append(id_el)

    checkbox = OxmlElement('w14:checkbox')
    checked = OxmlElement('w14:checked')
    checked.set(qn('w14:val'), '0')
    checkedState = OxmlElement('w14:checkedState')
    checkedState.set(qn('w14:val'), '2612')
    checkedState.set(qn('w14:font'), 'MS Gothic')
    uncheckedState = OxmlElement('w14:uncheckedState')
    uncheckedState.set(qn('w14:val'), '2610')
    uncheckedState.set(qn('w14:font'), 'MS Gothic')
    checkbox.append(checked)
    checkbox.append(checkedState)
    checkbox.append(uncheckedState)
    sdtPr.append(checkbox)
    sdt.append(sdtPr)

    sdtContent = OxmlElement('w:sdtContent')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'MS Gothic')
    rFonts.set(qn('w:hAnsi'), 'MS Gothic')
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = '\u2610'
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)

    p._p.append(sdt)
    run = p.add_run('  ' + text)
    run.font.size = Pt(style_size)
    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def _add_section_heading(doc, text):
    h = doc.add_heading(text, level=2)
    return h


def render_report_docx(report):
    """report: the structured dict returned by Claude (see PROMPT.md schema)."""
    doc = Document()

    title = doc.add_heading('🎄 Yule Love Lights', level=0)
    subtitle = doc.add_heading(
        f"Ops Report — {report['report_type']}", level=1
    )
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{report['date_range']}  •  {report['session_count']} session(s)"
    )
    meta_run.italic = True

    if report.get('data_quality_note'):
        dq = doc.add_paragraph()
        dq_run = dq.add_run(report['data_quality_note'])
        dq_run.italic = True
        dq_run.font.size = Pt(9)
        dq_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # Master Action List
    doc.add_heading('Master Action List (Open Items)', level=1)
    owners_order = ['Naldo', 'Jason', 'Both', 'TJ_Social']
    owner_labels = {'Naldo': 'Naldo', 'Jason': 'Jason', 'Both': 'Both / Decide Together', 'TJ_Social': 'Naldo — TJ / Social Media'}
    action_list = report.get('master_action_list', {})
    for owner in owners_order:
        items = action_list.get(owner, [])
        if not items:
            continue
        _add_section_heading(doc, owner_labels[owner])
        for item in items:
            line = item['item']
            if item.get('date_raised'):
                line += f"  [{item['date_raised']}]"
            if item.get('age_note'):
                line += f"  ⚠ {item['age_note']}"
            _add_checkbox_paragraph(doc, line)

    # Closed since last report
    closed = report.get('closed_since_last_report', [])
    if closed:
        doc.add_heading('Closed Since Last Report', level=1)
        for c in closed:
            _add_bullet(doc, c)

    # Notable numbers
    numbers = report.get('notable_numbers', [])
    if numbers:
        doc.add_heading('Notable Numbers', level=1)
        for n in numbers:
            _add_bullet(doc, n)

    # Decisions / notes
    decisions = report.get('decisions_notes', [])
    if decisions:
        doc.add_heading('Decisions / Notes', level=1)
        for d in decisions:
            _add_bullet(doc, d)

    # Important conversations to have
    conversations = report.get('conversations_to_have', [])
    if conversations:
        doc.add_heading('Important Conversations To Have', level=1)
        for c in conversations:
            _add_bullet(doc, c)

    # Day-by-day (Daily and Weekly only)
    day_by_day = report.get('day_by_day', [])
    if day_by_day:
        doc.add_heading('Day-by-Day Detail', level=1)
        for day in day_by_day:
            doc.add_heading(f"{day['date']} ({day.get('sessions', '?')} session(s))", level=2)
            if day.get('action_items'):
                p = doc.add_paragraph()
                p.add_run('Action items').italic = True
                for a in day['action_items']:
                    _add_bullet(doc, a)
            if day.get('decisions_notes'):
                p = doc.add_paragraph()
                p.add_run('Decisions / notes').italic = True
                for d in day['decisions_notes']:
                    _add_bullet(doc, d)
            for meeting in day.get('meetings', []):
                doc.add_heading(meeting['title'], level=3)
                for note in meeting.get('notes', []):
                    _add_bullet(doc, note)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
